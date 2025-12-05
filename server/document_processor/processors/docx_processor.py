import os
from typing import Dict, List
import logging

logger = logging.getLogger(__name__)

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.error("python-docx not available")

from . import BaseProcessor


class DOCXProcessor(BaseProcessor):
    def process(self, file_path: str) -> Dict:
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed")
        
        logger.info(f"Processing DOCX: {file_path}")
        
        result = {
            'text': '',
            'tables': [],
            'images': [],
            'formulas': [],
            'metadata': {}
        }
        
        try:
            doc = Document(file_path)
            
            result['text'] = self._extract_text(doc)
            result['tables'] = self._extract_tables(doc)
            result['metadata'] = self._extract_metadata(doc, file_path)
            
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            raise
        
        logger.info(f"DOCX processing complete: {len(result['text'])} chars, {len(result['tables'])} tables")
        
        return result
    
    def _extract_text(self, doc) -> str:
        text_parts = []
        
        for element in doc.element.body:
            if element.tag.endswith('p'):
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break
                
                if para:
                    text = para.text.strip()
                    
                    if text:
                        style = para.style.name if para.style else 'Normal'
                        
                        if 'Heading' in style:
                            level = int(style.split()[-1]) if any(c.isdigit() for c in style) else 1
                            text_parts.append(f"{'#' * level} {text}")
                        else:
                            text_parts.append(text)
            
            elif element.tag.endswith('tbl'):
                text_parts.append("[TABLE]")
        
        return '\n\n'.join(text_parts)
    
    def _extract_tables(self, doc) -> List[Dict]:
        tables = []
        
        for table_index, table in enumerate(doc.tables):
            table_dict = {
                'id': f'table_{table_index:03d}',
                'data': [],
                'shape': (len(table.rows), len(table.columns))
            }
            
            for row_index, row in enumerate(table.rows):
                row_data = []
                
                for cell in row.cells:
                    cell_text = ''.join([p.text for p in cell.paragraphs])
                    row_data.append(cell_text)
                
                table_dict['data'].append(row_data)
            
            tables.append(table_dict)
        
        logger.info(f"Extracted {len(tables)} tables")
        return tables
    
    def _extract_metadata(self, doc, file_path: str) -> Dict:
        metadata = {
            'file_size': os.path.getsize(file_path),
            'file_name': os.path.basename(file_path),
            'total_paragraphs': len(doc.paragraphs),
            'total_tables': len(doc.tables),
        }
        
        try:
            core_props = doc.core_properties
            
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.subject:
                metadata['subject'] = core_props.subject
            if core_props.created:
                metadata['created'] = str(core_props.created)
            if core_props.modified:
                metadata['modified'] = str(core_props.modified)
                
        except Exception as e:
            logger.warning(f"Core properties extraction failed: {e}")
        
        return metadata
